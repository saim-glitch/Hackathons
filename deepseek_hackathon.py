import os
import json
import ast
import requests
import subprocess
import streamlit as st
from tempfile import NamedTemporaryFile
from fpdf import FPDF
from docx import Document
from openai import OpenAI

# AIML API Configuration
base_url = "https://api.aimlapi.com/v1"
api_key = "9dc265c2a5204831b7fd062d4b20b5f0"  # Replace with your actual AIML API key
system_prompt = "You are a code reviewer. Analyze the code for vulnerabilities, efficiency, and best practices."
api = OpenAI(api_key=api_key, base_url=base_url)

# Function to analyze code using AIML API
def analyze_with_aiml(code: str):
    user_prompt = f"Analyze the following Python code for vulnerabilities, efficiency, and best practices:\n\n{code}"
    try:
        completion = api.chat.completions.create(
            model="deepseek/deepseek-r1",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.7,
            max_tokens=256,
        )
        return completion.choices[0].message.content
    except Exception as e:
        return {"error": f"AIML API request failed: {e}"}

# Function to perform static analysis using Bandit
def analyze_with_bandit(file_path: str):
    try:
        result = subprocess.run(["bandit", "-r", file_path, "-f", "json"], capture_output=True, text=True, check=True)
        return json.loads(result.stdout) if result.stdout else {}
    except subprocess.CalledProcessError as e:
        return {"error": f"Bandit analysis failed: {e}"}
    except FileNotFoundError:
        return {"error": "Bandit not found. Install it using `pip install bandit`"}

# Function to analyze code complexity using AST
def analyze_code_complexity(code: str):
    try:
        tree = ast.parse(code)
        functions = [node.name for node in ast.walk(tree) if isinstance(node, ast.FunctionDef)]
        classes = [node.name for node in ast.walk(tree) if isinstance(node, ast.ClassDef)]
        variables = [node.targets[0].id for node in ast.walk(tree) if isinstance(node, ast.Assign) and isinstance(node.targets[0], ast.Name)]
        return {
            "total_functions": len(functions),
            "total_classes": len(classes),
            "total_variables": len(variables),
            "functions": functions,
            "classes": classes,
            "variables": variables,
            "complexity_score": len(functions) * 1.5
        }
    except Exception as e:
        return {"error": str(e)}

# Function to check dependencies for security vulnerabilities
def check_dependencies():
    try:
        result = subprocess.run(["pip-audit", "--format", "json"], capture_output=True, text=True, check=True)
        return json.loads(result.stdout) if result.stdout else {}
    except subprocess.CalledProcessError as e:
        return {"error": f"Dependency check failed: {e}"}
    except FileNotFoundError:
        return {"error": "pip-audit not found. Install it using `pip install pip-audit`"}

# Function to generate a structured report
def generate_report(file_path: str, aiml_results, bandit_results, complexity_results, dependencies_results):
    report = {
        "File Analyzed": os.path.basename(file_path),
        "AIML Analysis": aiml_results,
        "Security Analysis (Bandit)": bandit_results,
        "Code Complexity": complexity_results,
        "Dependency Analysis": dependencies_results,
        "Recommendations": generate_recommendations(aiml_results, bandit_results, complexity_results, dependencies_results)
    }
    return report

# Function to generate recommendations based on analysis
def generate_recommendations(aiml, bandit, complexity, dependencies):
    recommendations = []
    
    if "issues" in bandit and bandit["issues"]:
        recommendations.append("Your code has security vulnerabilities. Review Bandit's findings and apply recommended fixes.")
    if "complexity_score" in complexity and complexity["complexity_score"] > 5:
        recommendations.append("Your code is quite complex. Consider refactoring long functions into smaller, reusable ones.")
    if "error" in dependencies:
        recommendations.append("Dependency security check failed. Ensure `pip-audit` is installed and run manually.")
    
    return recommendations if recommendations else ["No critical issues detected. Your code looks good!"]

# PDF Report Generation - Modified to output plain text instead of JSON
def generate_pdf_report(report, file_path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, txt=f"Code Analysis Report: {os.path.basename(file_path)}", ln=True, align="C")
    pdf.ln(10)
    
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 10, f"File Analyzed: {os.path.basename(file_path)}\n")
    pdf.multi_cell(0, 10, "AIML Analysis:\n" + str(report['AIML Analysis']) + "\n")
    
    # Security Analysis - Format as plain text
    sec_text = "Security Analysis (Bandit):\n"
    if "error" in report["Security Analysis (Bandit)"]:
        sec_text += report["Security Analysis (Bandit)"]["error"] + "\n"
    elif report["Security Analysis (Bandit)"].get("results"):
        for result in report["Security Analysis (Bandit)"]["results"]:
            sec_text += f"Test: {result.get('test_name', 'N/A')}\n"
            sec_text += f"Severity: {result.get('issue_severity', 'N/A')}\n"
            sec_text += f"Confidence: {result.get('issue_confidence', 'N/A')}\n"
            sec_text += f"Code: {result.get('code', '').strip()}\n"
            sec_text += f"Explanation: {result.get('issue_text', 'N/A')}\n"
            sec_text += "--------------------------------------\n"
    else:
        sec_text += "No security issues found.\n"
    pdf.multi_cell(0, 10, sec_text + "\n")
    
    # Code Complexity - Format as plain text
    cc = report["Code Complexity"]
    cc_text = ("Code Complexity:\n" +
               f"Total Functions: {cc.get('total_functions', 0)}\n" +
               f"Total Classes: {cc.get('total_classes', 0)}\n" +
               f"Total Variables: {cc.get('total_variables', 0)}\n" +
               f"Complexity Score: {cc.get('complexity_score', 0)}\n" +
               f"Functions: {', '.join(cc.get('functions', []))}\n" +
               f"Classes: {', '.join(cc.get('classes', []))}\n" +
               f"Variables: {', '.join(cc.get('variables', []))}\n")
    pdf.multi_cell(0, 10, cc_text + "\n")
    
    # Dependency Analysis - Format as plain text
    dep = report["Dependency Analysis"]
    dep_text = "Dependency Analysis:\n"
    if "error" in dep:
        dep_text += dep["error"] + "\n"
    elif dep.get("vulnerabilities"):
        for vuln in dep["vulnerabilities"]:
            dep_text += f"Name: {vuln.get('name', 'N/A')}\n"
            dep_text += f"Spec: {vuln.get('spec', 'N/A')}\n"
            advisory = vuln.get("advisory", {})
            dep_text += f"Severity: {advisory.get('severity', 'N/A')}\n"
            dep_text += f"Affected Versions: {advisory.get('affected_versions', 'N/A')}\n"
            dep_text += f"Description: {advisory.get('description', 'N/A')}\n"
            dep_text += "--------------------------------------\n"
    else:
        dep_text += "No dependency vulnerabilities found.\n"
    pdf.multi_cell(0, 10, dep_text + "\n")
    
    # Recommendations
    rec_text = "Recommendations:\n" + "\n".join(report["Recommendations"]) + "\n"
    pdf.multi_cell(0, 10, rec_text)
    
    pdf_output = f"{file_path}_report.pdf"
    pdf.output(pdf_output)
    return pdf_output

# Word Report Generation - Modified to output plain text instead of JSON
def generate_word_report(report, file_path):
    doc = Document()
    doc.add_heading(f"Code Analysis Report: {os.path.basename(file_path)}", 0)
    
    doc.add_heading("File Analyzed", level=1)
    doc.add_paragraph(f"{os.path.basename(file_path)}")
    
    doc.add_heading("AIML Analysis", level=1)
    doc.add_paragraph(str(report['AIML Analysis']))
    
    # Security Analysis - Format as plain text
    sec_text = "Security Analysis (Bandit):\n"
    if "error" in report["Security Analysis (Bandit)"]:
        sec_text += report["Security Analysis (Bandit)"]["error"] + "\n"
    elif report["Security Analysis (Bandit)"].get("results"):
        for result in report["Security Analysis (Bandit)"]["results"]:
            sec_text += f"Test: {result.get('test_name', 'N/A')}\n"
            sec_text += f"Severity: {result.get('issue_severity', 'N/A')}\n"
            sec_text += f"Confidence: {result.get('issue_confidence', 'N/A')}\n"
            sec_text += f"Code: {result.get('code', '').strip()}\n"
            sec_text += f"Explanation: {result.get('issue_text', 'N/A')}\n"
            sec_text += "--------------------------------------\n"
    else:
        sec_text += "No security issues found.\n"
    doc.add_heading("Security Analysis (Bandit)", level=1)
    doc.add_paragraph(sec_text)
    
    # Code Complexity - Format as plain text
    cc = report["Code Complexity"]
    cc_text = ("Code Complexity:\n" +
               f"Total Functions: {cc.get('total_functions', 0)}\n" +
               f"Total Classes: {cc.get('total_classes', 0)}\n" +
               f"Total Variables: {cc.get('total_variables', 0)}\n" +
               f"Complexity Score: {cc.get('complexity_score', 0)}\n" +
               f"Functions: {', '.join(cc.get('functions', []))}\n" +
               f"Classes: {', '.join(cc.get('classes', []))}\n" +
               f"Variables: {', '.join(cc.get('variables', []))}\n")
    doc.add_heading("Code Complexity", level=1)
    doc.add_paragraph(cc_text)
    
    # Dependency Analysis - Format as plain text
    dep = report["Dependency Analysis"]
    dep_text = "Dependency Analysis:\n"
    if "error" in dep:
        dep_text += dep["error"] + "\n"
    elif dep.get("vulnerabilities"):
        for vuln in dep["vulnerabilities"]:
            dep_text += f"Name: {vuln.get('name', 'N/A')}\n"
            dep_text += f"Spec: {vuln.get('spec', 'N/A')}\n"
            advisory = vuln.get("advisory", {})
            dep_text += f"Severity: {advisory.get('severity', 'N/A')}\n"
            dep_text += f"Affected Versions: {advisory.get('affected_versions', 'N/A')}\n"
            dep_text += f"Description: {advisory.get('description', 'N/A')}\n"
            dep_text += "--------------------------------------\n"
    else:
        dep_text += "No dependency vulnerabilities found.\n"
    doc.add_heading("Dependency Analysis", level=1)
    doc.add_paragraph(dep_text)
    
    # Recommendations
    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph("\n".join(report["Recommendations"]))
    
    word_output = f"{file_path}_report.docx"
    doc.save(word_output)
    return word_output

def main():
    st.title("🚀 AI-Powered Code Review & Vulnerability Detector")
    st.markdown("Upload your Python file to analyze for vulnerabilities, efficiency, and best practices.")

    uploaded_file = st.file_uploader("📂 Upload a Python file", type=["py"])
    if uploaded_file:
        with NamedTemporaryFile(delete=False, suffix=".py") as temp_file:
            temp_file.write(uploaded_file.read())
            temp_file_path = temp_file.name
        
        with open(temp_file_path, "r") as f:
            code = f.read()
        
        st.subheader("Uploaded Code:")
        st.code(code, language="python")
        
        # Initialize analysis variables with default values
        aiml_results = None
        bandit_results = None
        complexity_results = None
        dependencies_results = None
        
        try:
            with st.spinner("Analyzing code... Please wait."):
                aiml_results = analyze_with_aiml(code)
                bandit_results = analyze_with_bandit(temp_file_path)
                complexity_results = analyze_code_complexity(code)
                dependencies_results = check_dependencies()
        except Exception as e:
            st.error(f"An error occurred during analysis: {str(e)}")
            return  # Exit the function if analysis fails
        
        # Generate the report
        report = generate_report(temp_file_path, aiml_results, bandit_results, complexity_results, dependencies_results)
        
        # Display results only if analysis was successful
        if aiml_results is not None:
            st.subheader("Analysis Results")
            
            # Section 1: AI Analysis
            st.markdown("---")
            st.header("🧠 AI Code Review Findings")
            st.markdown("#### Code Quality Assessment")
            st.write(aiml_results)

            # Section 2: Security Analysis
            st.markdown("---")
            st.header("🔒 Security Analysis Results")
            if "error" in bandit_results:
                st.error(bandit_results["error"])
            else:
                if bandit_results.get("results"):
                    for result in bandit_results["results"]:
                        with st.expander(f"🔍 {result['test_name']}", expanded=False):
                            st.markdown(f"**Severity:** {result['issue_severity']}")
                            st.markdown(f"**Confidence:** {result['issue_confidence']}")
                            st.code(result['code'], language='python')
                            st.markdown(f"**Explanation:** {result['issue_text']}")
                else:
                    st.success("✅ No security issues found by Bandit")

            # Section 3: Code Complexity
            st.markdown("---")
            st.header("📊 Code Complexity Analysis")
            cols = st.columns(3)
            cols[0].metric("Functions", complexity_results["total_functions"])
            cols[1].metric("Classes", complexity_results["total_classes"])
            cols[2].metric("Complexity Score", complexity_results["complexity_score"])

            st.markdown("#### Identified Elements")
            st.write(f"**Functions:** {', '.join(complexity_results['functions']) or 'None'}")
            st.write(f"**Classes:** {', '.join(complexity_results['classes']) or 'None'}")
            st.write(f"**Variables:** {', '.join(complexity_results['variables']) or 'None'}")

            # Section 4: Dependency Analysis
            st.markdown("---")
            st.header("📦 Dependency Vulnerabilities")
            if "error" in dependencies_results:
                st.error(dependencies_results["error"])
            else:
                if dependencies_results.get("vulnerabilities"):
                    for vuln in dependencies_results["vulnerabilities"]:
                        st.markdown(f"### {vuln['name']} ({vuln['spec']})")
                        st.markdown(f"**Severity:** `{vuln['advisory']['severity']}`")
                        st.markdown(f"**Affected Versions:** {vuln['advisory']['affected_versions']}")
                        st.markdown(f"**Description:** {vuln['advisory']['description']}")
                else:
                    st.success("✅ All dependencies are secure")

            # Section 5: Recommendations
            st.markdown("---")
            st.header("📝 Key Recommendations")
            for rec in report["Recommendations"]:
                if "vulnerabilities" in rec:
                    st.error(f"🚨 {rec}")
                elif "complex" in rec:
                    st.warning(f"⚠️ {rec}")
                else:
                    st.success(f"✅ {rec}")

            # Download buttons for PDF and Word report
            pdf_file = generate_pdf_report(report, temp_file_path)
            word_file = generate_word_report(report, temp_file_path)
            
            st.download_button("Download PDF Report", open(pdf_file, "rb").read(), file_name=f"{os.path.basename(temp_file_path)}_report.pdf", mime="application/pdf")
            st.download_button("Download Word Report", open(word_file, "rb").read(), file_name=f"{os.path.basename(temp_file_path)}_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            # Clean up temporary files
            os.remove(temp_file_path)
            os.remove(pdf_file)
            os.remove(word_file)
        else:
            st.error("Analysis failed. Please check the uploaded file and try again.")

if __name__ == "__main__":
    main()
